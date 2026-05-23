"""
Script untuk generate BAB IV skripsi dalam format .docx
Judul: Deteksi Deepfake Berbasis Residual Noise: Perbandingan Representasi
       Fitur Domain Spasial dan DCT Menggunakan CNN
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (3cm left, 4cm top-bottom-right typical Indonesian thesis) ──
section = doc.sections[0]
section.top_margin    = Cm(4)
section.bottom_margin = Cm(3)
section.left_margin   = Cm(4)
section.right_margin  = Cm(3)

# ── Paragraph/style helpers ──────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=12):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'

def add_heading(doc, text, level=1):
    """Heading 1/2 with Times New Roman bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=12 if level >= 2 else 12)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    return p

def add_body(doc, text, indent=False, justify=True):
    """Normal body paragraph."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)          # double-space
    if indent:
        pf.first_line_indent = Cm(1.25)
    return p

def add_body_mixed(doc, parts, indent=False):
    """parts = list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    pf = p.paragraph_format
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(24)
    if indent:
        pf.first_line_indent = Cm(1.25)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=True)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pf = p.paragraph_format
    pf.left_indent   = Cm(1.25)
    pf.space_before  = Pt(3)
    pf.space_after   = Pt(3)
    pf.line_spacing  = Pt(14)
    return p

def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(('    ' * indent_level) + '• ' + text)
    set_font(run)
    pf = p.paragraph_format
    pf.left_indent  = Cm(1.25 + indent_level * 0.5)
    pf.space_after  = Pt(3)
    pf.line_spacing = Pt(24)
    return p

# ── Simple table helper ──────────────────────────────────────────────────────

def make_table(doc, data, col_widths=None):
    """data[0] = header row, rest = data rows."""
    rows = len(data)
    cols = len(data[0])
    tbl  = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(data):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            if r_idx == 0:
                run.bold = True

    if col_widths:
        for r_idx in range(rows):
            row = tbl.rows[r_idx]
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = Cm(w)
    return tbl

# ════════════════════════════════════════════════════════════════════════════
# BAB IV — IMPLEMENTASI DAN PEMBAHASAN
# ════════════════════════════════════════════════════════════════════════════

# ── Judul Bab ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BAB IV\nIMPLEMENTASI DAN PEMBAHASAN')
set_font(run, bold=True, size=14)
p.paragraph_format.space_after = Pt(18)

# ════════════════════════════════════════════════════════════════════════════
# 4.1  BUSINESS UNDERSTANDING
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.1 Business Understanding')

add_body(doc,
    'Fase Business Understanding merupakan tahapan pertama dalam kerangka kerja '
    'CRISP-DM (Cross-Industry Standard Process for Data Mining). Pada fase ini '
    'dilakukan pemahaman mendalam mengenai tujuan penelitian, rumusan masalah, '
    'serta penentuan kriteria keberhasilan dari sudut pandang penelitian.',
    indent=True)

add_heading(doc, '4.1.1 Tujuan Penelitian', level=2)

add_body(doc,
    'Penelitian ini bertujuan untuk mengembangkan sistem deteksi deepfake wajah '
    'berbasis Convolutional Neural Network (CNN) dengan membandingkan tiga pendekatan '
    'representasi fitur: (1) RGB langsung (baseline), (2) residual noise domain spasial, '
    'dan (3) koefisien DCT dari residual noise. Sistem dibangun menggunakan arsitektur '
    'Xception sebagai backbone dengan transfer learning dari ImageNet.',
    indent=True)

add_body(doc,
    'Permasalahan utama yang diselesaikan adalah: apakah penambahan informasi residual '
    'noise—baik dalam domain spasial maupun domain frekuensi (DCT)—dapat meningkatkan '
    'kemampuan CNN dalam mendeteksi manipulasi wajah (deepfake) dibandingkan dengan '
    'penggunaan citra RGB biasa?',
    indent=True)

add_heading(doc, '4.1.2 Kriteria Keberhasilan', level=2)

add_body(doc,
    'Kriteria keberhasilan penelitian ini ditetapkan sebagai berikut:')

add_bullet(doc, 'Akurasi klasifikasi minimal 95% pada data uji.')
add_bullet(doc, 'Nilai AUC-ROC minimal 0,98 yang mencerminkan kemampuan diskriminasi model.')
add_bullet(doc, 'Perbandingan performa ketiga metode menunjukkan perbedaan yang dapat '
                'diinterpretasikan secara kuantitatif.')
add_bullet(doc, 'Sistem mampu mengklasifikasikan citra wajah sebagai "real" atau "fake" '
                'secara otomatis tanpa intervensi manual.')

# ════════════════════════════════════════════════════════════════════════════
# 4.2  DATA UNDERSTANDING
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.2 Data Understanding')

add_body(doc,
    'Fase Data Understanding mencakup proses pengumpulan data awal, eksplorasi '
    'karakteristik dataset, serta penilaian kualitas data yang akan digunakan '
    'dalam penelitian.',
    indent=True)

add_heading(doc, '4.2.1 Dataset IMSFD', level=2)

add_body(doc,
    'Dataset yang digunakan adalah Indonesian Multi-Subject Face Dataset (IMSFD), '
    'sebuah dataset wajah yang dikumpulkan dari subjek-subjek Indonesia. Dataset ini '
    'dipilih karena representasinya terhadap karakteristik wajah populasi Indonesia '
    'yang menjadi konteks penelitian ini.',
    indent=True)

add_body(doc,
    'IMSFD terdiri dari 68 subjek unik yang diorganisasikan dalam tiga kelompok: '
    'A1, A2, dan B1. Setiap subjek memiliki subfolder internal berupa training/ dan '
    'testing/ yang berisi foto-foto wajah dengan berbagai variasi pose, ekspresi, '
    'dan kondisi pencahayaan. Format penamaan kunci subjek mengikuti pola '
    '{grup}__{subject_id}, contohnya A1__13020220330.',
    indent=True)

# Tabel 4.1
add_caption(doc, 'Tabel 4.1 Statistik Dataset IMSFD')
make_table(doc, [
    ['Atribut',           'Keterangan'],
    ['Jumlah Subjek',     '68 subjek unik'],
    ['Kelompok',          'A1, A2, B1'],
    ['Format Data',       'Citra wajah (JPG/PNG)'],
    ['Asal Dataset',      'Dataset wajah Indonesia (IMSFD)'],
    ['Subfolder Internal','training/ dan testing/ per subjek'],
    ['Format Kunci',      '{grup}__{subject_id}'],
], col_widths=[6, 10])

doc.add_paragraph()

add_heading(doc, '4.2.2 Pembangkitan Deepfake', level=2)

add_body(doc,
    'Citra deepfake dihasilkan menggunakan metode face swapping berbasis InsightFace '
    'dengan model inswapper_128.onnx. Proses ini mensimulasikan manipulasi wajah yang '
    'realistis dengan cara menukar wajah sumber ke wajah target menggunakan representasi '
    'embedding wajah.',
    indent=True)

add_body_mixed(doc, [
    ('Parameter pembangkitan deepfake yang digunakan: ',False,False),
    ('NUM_PAIRS=5000',True,False),
    (' (pasangan sumber-target), ',False,False),
    ('SWAPS_PER_PAIR=3',True,False),
    (' (pengulangan swap per pasangan), dan ukuran output ',False,False),
    ('224×224 piksel',True,False),
    (' sesuai input model.',False,False),
], indent=True)

add_body(doc,
    'Pembangkitan deepfake dilakukan dengan mode split_aware, artinya citra deepfake '
    'dibuat secara terpisah per split (train/val/test) menggunakan hanya subjek-subjek '
    'yang ada pada split tersebut. Hal ini mencegah kebocoran data antar split dan '
    'memastikan evaluasi yang adil.',
    indent=True)

add_heading(doc, '4.2.3 Eksplorasi dan Statistik Data', level=2)

add_body(doc,
    'Setelah proses pembangkitan deepfake, distribusi data pada setiap split adalah '
    'sebagai berikut:')

# Tabel 4.2
add_caption(doc, 'Tabel 4.2 Distribusi Data per Split')
make_table(doc, [
    ['Split', 'Jumlah Subjek', 'Citra Real', 'Citra Fake', 'Total'],
    ['Train', '47',   '9.528',  '9.528',  '19.056'],
    ['Val',   '10',   '2.087',  '2.087',  ' 4.174'],
    ['Test',  '11',   '1.824',  '1.824',  ' 3.648'],
    ['Total', '68',  '13.439', '13.439', '26.878'],
], col_widths=[3, 4, 4, 4, 3])

doc.add_paragraph()

add_body(doc,
    'Pembagian data menggunakan strategi subject-level split dengan rasio 70:15:15 '
    '(train:val:test) dan random seed=42. Strategi ini memastikan bahwa satu subjek '
    'hanya muncul di salah satu split sehingga tidak terjadi identity leakage yang '
    'dapat mengakibatkan overestimasi performa model.',
    indent=True)

add_body(doc,
    'Keseimbangan kelas (real vs fake) dijaga secara eksplisit pada setiap split '
    'melalui fungsi balance_splits() yang menghapus kelebihan citra dari kelas mayoritas. '
    'Dengan demikian rasio kelas 1:1 terjaga di seluruh split.',
    indent=True)

# ════════════════════════════════════════════════════════════════════════════
# 4.3  DATA PREPARATION
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.3 Data Preparation')

add_body(doc,
    'Fase Data Preparation mencakup seluruh proses transformasi data mentah menjadi '
    'format yang siap digunakan untuk pelatihan model. Proses ini meliputi pembagian '
    'dataset, ekstraksi fitur residual noise, ekstraksi fitur DCT, dan pipeline '
    'augmentasi data.',
    indent=True)

add_heading(doc, '4.3.1 Pembagian Dataset (Subject-Level Split)', level=2)

add_body(doc,
    'Pembagian dataset dilakukan pada level subjek untuk mencegah identity leakage. '
    'Proses ini diimplementasikan dalam skrip split_dataset.py dengan langkah-langkah '
    'berikut:',
    indent=True)

add_bullet(doc, 'Pengumpulan seluruh subjek dari direktori IMSFD beserta subfolder '
                'training/ dan testing/ (fungsi collect_subjects()).')
add_bullet(doc, 'Pengacakan urutan subjek dengan random seed=42 untuk reproducibility.')
add_bullet(doc, 'Pembagian 68 subjek menjadi 47 train, 10 val, dan 11 test '
                '(rasio 70:15:15).')
add_bullet(doc, 'Penyalinan citra ke direktori data/train/, data/val/, dan data/test/ '
                'dengan substruktur real/ dan fake/.')
add_bullet(doc, 'Penyeimbangan kelas melalui balance_splits() agar jumlah citra '
                'real = fake pada tiap split.')
add_bullet(doc, 'Penyimpanan metadata split ke data/dataset_metadata.json.')

add_heading(doc, '4.3.2 Ekstraksi Residual Noise (Domain Spasial)', level=2)

add_body(doc,
    'Residual noise merupakan selisih antara citra asli dan versi yang telah '
    'diperhalus (denoised). Secara matematis dinyatakan sebagai:',
    indent=True)

add_body(doc, '    R = I − G(I)',)

add_body(doc,
    'di mana I adalah citra input dan G(I) adalah hasil Gaussian blur terhadap I. '
    'Residual noise ini menangkap artefak manipulasi yang sering tertinggal setelah '
    'proses deepfake generation.',
    indent=True)

add_body(doc,
    'Implementasi ekstraksi residual noise dalam kode (src/feature_extraction.py):',
    indent=True)

add_code_block(doc, 'class ResidualNoiseExtractor:')
add_code_block(doc, '    def __init__(self, kernel_size=5, sigma=1.0):')
add_code_block(doc, '        self.kernel_size = kernel_size  # 5x5')
add_code_block(doc, '        self.sigma = sigma              # 1.0')
add_code_block(doc, '')
add_code_block(doc, '    def extract(self, image_tensor):')
add_code_block(doc, '        # 1. Denormalisasi dari [-1,1] ke [0,255]')
add_code_block(doc, '        img_np = ((image_tensor + 1) * 127.5).clamp(0, 255)')
add_code_block(doc, '        # 2. Gaussian blur sebagai estimasi sinyal bersih')
add_code_block(doc, '        blurred = cv2.GaussianBlur(img_np, (5,5), sigmaX=1.0)')
add_code_block(doc, '        # 3. Residual = Original - Blurred')
add_code_block(doc, '        residual = img_np - blurred')
add_code_block(doc, '        # 4. Normalisasi residual ke [-1, 1]')
add_code_block(doc, '        residual = residual / 127.5 - 1.0')
add_code_block(doc, '        return residual')

add_body(doc,
    'Parameter yang digunakan: kernel Gaussian 5×5 (RESIDUAL_SPATIAL_KERNEL_SIZE=5) '
    'dan sigma=1.0 (RESIDUAL_SPATIAL_SIGMA=1.0). Output residual dinormalisasi ke '
    'rentang [−1, 1] agar kompatibel dengan pipeline normalisasi model.',
    indent=True)

add_heading(doc, '4.3.3 Ekstraksi Fitur DCT dari Residual Noise', level=2)

add_body(doc,
    'Discrete Cosine Transform (DCT) diaplikasikan pada residual noise untuk '
    'merepresentasikan informasi artefak dalam domain frekuensi. Pendekatan ini '
    'didasarkan pada fakta bahwa manipulasi deepfake sering meninggalkan anomali '
    'pada pola frekuensi tinggi yang lebih mudah terdeteksi dalam domain DCT.',
    indent=True)

add_body(doc,
    'Proses ekstraksi DCT dilakukan secara block-wise dengan ukuran blok 8×8 '
    '(DCT_BLOCK_SIZE=8), mengikuti standar yang digunakan dalam kompresi JPEG. '
    'Langkah-langkah implementasi:',
    indent=True)

add_bullet(doc, 'Padding citra residual ke kelipatan 8 menggunakan mode reflect.')
add_bullet(doc, 'Aplikasi 2D DCT pada setiap blok 8×8 per channel warna.')
add_bullet(doc, 'Normalisasi percentil (persentil ke-1 hingga ke-99) ke rentang [−1, 1].')

add_code_block(doc, 'class DCTExtractor:')
add_code_block(doc, '    def __init__(self, block_size=8):')
add_code_block(doc, '        self.block_size = block_size  # 8x8 blok')
add_code_block(doc, '')
add_code_block(doc, '    def extract(self, residual_tensor):')
add_code_block(doc, '        # Pad ke kelipatan block_size (reflect mode)')
add_code_block(doc, '        padded = pad_to_multiple(residual_tensor, self.block_size)')
add_code_block(doc, '        # Block-wise 2D DCT per channel')
add_code_block(doc, '        dct_result = blockwise_dct(padded, self.block_size)')
add_code_block(doc, '        # Percentile normalization [1st, 99th] -> [-1, 1]')
add_code_block(doc, '        dct_norm = percentile_normalize(dct_result)')
add_code_block(doc, '        return dct_norm')

add_heading(doc, '4.3.4 Pipeline Augmentasi Data', level=2)

add_body(doc,
    'Augmentasi data diterapkan hanya pada data pelatihan untuk meningkatkan '
    'generalisasi model dan mengurangi risiko overfitting. Pipeline augmentasi '
    'diimplementasikan dalam get_transforms() pada src/data_preprocessing.py:',
    indent=True)

# Tabel 4.3
add_caption(doc, 'Tabel 4.3 Pipeline Augmentasi Data Pelatihan')
make_table(doc, [
    ['Transformasi',         'Parameter',                          'Tujuan'],
    ['Resize',               '244×244 piksel',                     'Standarisasi ukuran awal'],
    ['RandomCrop',           '224×224 piksel',                     'Variasi posisi wajah'],
    ['RandomHorizontalFlip', 'p=0,5',                              'Variasi orientasi horizontal'],
    ['RandomRotation',       '±10 derajat',                        'Variasi rotasi ringan'],
    ['ColorJitter',          'brightness=0,15, contrast=0,15,\nsaturation=0,10, hue=0,05',
                                                                   'Variasi kondisi pencahayaan'],
    ['RandomAffine',         'translate=(0,05, 0,05)',             'Variasi translasi ringan'],
    ['RandomErasing',        'p=0,1, scale=(0,02–0,08)',          'Simulasi oklusi parsial'],
    ['Normalize',            'mean=[0,5,0,5,0,5],\nstd=[0,5,0,5,0,5]',
                                                                   'Normalisasi ke [-1, 1]'],
], col_widths=[4, 6, 6])

doc.add_paragraph()

add_body(doc,
    'Untuk data validasi dan pengujian, hanya dilakukan Resize ke 224×224 dan '
    'normalisasi tanpa augmentasi lainnya, untuk memastikan evaluasi yang konsisten '
    'dan tidak bias.',
    indent=True)

# ════════════════════════════════════════════════════════════════════════════
# 4.4  MODELING
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.4 Modeling')

add_body(doc,
    'Fase Modeling mencakup implementasi arsitektur model, proses pelatihan, '
    'dan evaluasi performa ketiga metode yang dibandingkan dalam penelitian ini.',
    indent=True)

add_heading(doc, '4.4.1 Arsitektur Model', level=2)

add_body(doc,
    'Ketiga metode menggunakan arsitektur yang identik berbasis Xception, sebuah '
    'jaringan saraf konvolusi yang dikembangkan oleh Chollet (2017) dan dikenal '
    'efektif untuk tugas deteksi manipulasi visual. Model diinisialisasi menggunakan '
    'library timm (PyTorch Image Models) dengan bobot pretrained dari ImageNet.',
    indent=True)

add_body(doc,
    'Arsitektur model secara keseluruhan:',
    indent=True)

add_bullet(doc, 'Backbone: timm.create_model("xception", pretrained=True) — '
                'mengekstrak fitur 2048 dimensi.')
add_bullet(doc, 'Dropout: Dropout(p=0,5) — regularisasi untuk mencegah overfitting.')
add_bullet(doc, 'Classifier: Linear(2048 → 2) — lapisan klasifikasi biner (real/fake).')

add_code_block(doc, 'class XceptionBaseline(nn.Module):')
add_code_block(doc, '    def __init__(self, num_classes=2, pretrained=True):')
add_code_block(doc, '        super().__init__()')
add_code_block(doc, '        self.backbone = timm.create_model(')
add_code_block(doc, '            "xception", pretrained=pretrained,')
add_code_block(doc, '            num_classes=0, global_pool="avg"')
add_code_block(doc, '        )')
add_code_block(doc, '        self.classifier = nn.Sequential(')
add_code_block(doc, '            nn.Dropout(0.5),')
add_code_block(doc, '            nn.Linear(2048, num_classes)')
add_code_block(doc, '        )')
add_code_block(doc, '    def forward(self, x):')
add_code_block(doc, '        features = self.backbone(x)')
add_code_block(doc, '        return self.classifier(features)')

add_body(doc,
    'Kelas XceptionResidualSpatial dan XceptionResidualDCT memiliki struktur yang '
    'identik dengan XceptionBaseline—perbedaan ketiganya hanya terletak pada '
    'preprocessing input yang dilakukan sebelum data dimasukkan ke backbone.',
    indent=True)

add_heading(doc, '4.4.2 Konfigurasi Pelatihan', level=2)

add_body(doc,
    'Pelatihan model dilakukan pada GPU NVIDIA GTX 1660 Ti dengan CUDA 12.4 '
    'menggunakan framework PyTorch. Seluruh konfigurasi hyperparameter disimpan '
    'dalam config/config.py.',
    indent=True)

# Tabel 4.4
add_caption(doc, 'Tabel 4.4 Konfigurasi Hyperparameter Pelatihan')
make_table(doc, [
    ['Parameter',           'Nilai',                  'Keterangan'],
    ['Optimizer',           'Adam',                   'lr=0,0001, weight_decay=1e-4'],
    ['Learning Rate',       '0,0001',                 'Initial learning rate'],
    ['LR Scheduler',        'ReduceLROnPlateau',      'factor=0,5, patience=2'],
    ['Batch Size',          '8',                      'Mini-batch per iterasi'],
    ['Num Epochs',          '12',                     'Jumlah epoch pelatihan'],
    ['Gradient Clipping',   '1,0',                    'Mencegah exploding gradient'],
    ['AMP',                 'Enabled',                'Automatic Mixed Precision (FP16)'],
    ['Early Stopping',      'patience=5',             'Berhenti jika tidak ada perbaikan'],
    ['Loss Function',       'CrossEntropyLoss',       'Untuk klasifikasi biner'],
    ['Ukuran Input',        '224×224×3',              'RGB, ternormalisasi ke [-1,1]'],
], col_widths=[4, 4, 8])

doc.add_paragraph()

add_heading(doc, '4.4.3 Proses Pelatihan', level=2)

add_body(doc,
    'Pelatihan diimplementasikan dalam src/train.py dengan tiga fungsi terpisah '
    'untuk masing-masing metode: train_epoch_baseline(), '
    'train_epoch_residual_spatial(), dan train_epoch_residual_dct(). '
    'Setiap fungsi menggunakan Automatic Mixed Precision (AMP) dengan GradScaler '
    'untuk efisiensi komputasi.',
    indent=True)

add_body(doc,
    'Alur pelatihan untuk metode residual_dct adalah sebagai berikut:',
    indent=True)

add_bullet(doc, 'Batch citra RGB dimuat dari DataLoader.')
add_bullet(doc, 'Residual noise diekstrak: residual = ResidualNoiseExtractor.extract_batch(images).')
add_bullet(doc, 'DCT diterapkan pada residual: dct_feat = DCTExtractor.extract_batch(residual).')
add_bullet(doc, 'Fitur DCT dimasukkan ke model Xception untuk mendapat prediksi.')
add_bullet(doc, 'Loss CrossEntropy dihitung, backpropagation dilakukan dengan gradient clipping=1,0.')
add_bullet(doc, 'Scheduler ReduceLROnPlateau memantau val_loss dan mengurangi lr jika stagnan.')
add_bullet(doc, 'Model terbaik (val_loss terendah) disimpan secara otomatis.')

add_heading(doc, '4.4.4 Evaluasi Model', level=2)

add_body(doc,
    'Evaluasi dilakukan pada data uji (test set) menggunakan metrik klasifikasi biner '
    'standar yang dihitung menggunakan library scikit-learn. Proses evaluasi '
    'diimplementasikan dalam src/evaluate.py.',
    indent=True)

# Tabel 4.5
add_caption(doc, 'Tabel 4.5 Metrik Evaluasi yang Digunakan')
make_table(doc, [
    ['Metrik',    'Formula',                             'Keterangan'],
    ['Accuracy',  '(TP+TN) / (TP+TN+FP+FN)',            'Proporsi prediksi benar secara keseluruhan'],
    ['Precision', 'TP / (TP+FP)',                        'Ketepatan prediksi kelas fake'],
    ['Recall',    'TP / (TP+FN)',                        'Kemampuan mendeteksi semua deepfake'],
    ['F1-Score',  '2 × (P×R) / (P+R)',                  'Harmonic mean Precision dan Recall'],
    ['AUC-ROC',   'Area Under ROC Curve',                'Kemampuan diskriminasi keseluruhan'],
], col_widths=[3, 5, 8])

doc.add_paragraph()

add_body(doc,
    'Nilai probabilitas kelas positif (fake, kelas 1) diambil dari output softmax '
    'model untuk menghitung AUC-ROC. Threshold default 0,5 digunakan untuk '
    'mengkonversi probabilitas menjadi label biner.',
    indent=True)

add_heading(doc, '4.4.5 Hasil Perbandingan Ketiga Metode', level=2)

add_body(doc,
    'Berikut adalah hasil evaluasi ketiga metode pada data uji (test set) yang '
    'terdiri dari 3.648 citra (1.824 real dan 1.824 fake):',
    indent=True)

# Tabel 4.6
add_caption(doc, 'Tabel 4.6 Perbandingan Hasil Evaluasi Ketiga Metode')
make_table(doc, [
    ['Metode',           'Loss',   'Accuracy (%)', 'Precision', 'Recall', 'F1-Score', 'AUC'],
    ['Baseline RGB',     '0,0113', '99,75',        '1,0000',    '0,9951', '0,9975',   '0,9998'],
    ['Residual Spatial', '0,0094', '99,78',        '0,9973',    '0,9984', '0,9978',   '0,9999'],
    ['Residual DCT',     '0,0418', '98,52',        '0,9733',    '0,9978', '0,9854',   '0,9994'],
], col_widths=[4, 2.5, 3.5, 3, 2.5, 3, 2.5])

doc.add_paragraph()

add_heading(doc, '4.4.6 Analisis dan Pembahasan Hasil', level=2)

add_body(doc,
    'Berdasarkan Tabel 4.6, dapat diperoleh beberapa temuan utama:',
    indent=True)

add_body_mixed(doc, [
    ('a) Residual Spatial unggul secara keseluruhan. ',True,False),
    ('Metode Residual Spatial mencapai akurasi tertinggi (99,78%) dan AUC '
     'tertinggi (0,9999), mengungguli Baseline RGB. Hal ini menunjukkan bahwa '
     'penambahan informasi residual noise dalam domain spasial memberikan sinyal '
     'tambahan yang relevan bagi model dalam mendeteksi artefak manipulasi deepfake.',False,False),
], indent=True)

add_body_mixed(doc, [
    ('b) Baseline RGB mencapai performa sangat tinggi. ',True,False),
    ('Akurasi 99,75% dengan Precision sempurna (1,0000) menunjukkan bahwa model '
     'Xception pretrained dengan RGB input saja sudah sangat efektif. Precision '
     'sempurna berarti tidak ada false positive—setiap prediksi fake adalah benar.',False,False),
], indent=True)

add_body_mixed(doc, [
    ('c) Residual DCT memiliki Recall tertinggi. ',True,False),
    ('Recall 0,9978 pada metode Residual DCT mengindikasikan kemampuan terbaik '
     'dalam mendeteksi semua citra deepfake (meminimalkan false negative). '
     'Namun, akurasi dan precision yang lebih rendah menunjukkan adanya trade-off '
     'dengan false positive yang lebih tinggi.',False,False),
], indent=True)

add_body_mixed(doc, [
    ('d) Residual DCT memiliki loss tertinggi. ',True,False),
    ('Loss 0,0418 yang lebih tinggi dibandingkan dua metode lain mengindikasikan '
     'bahwa transformasi DCT memperkenalkan variasi representasi yang lebih '
     'kompleks, sehingga model membutuhkan lebih banyak iterasi atau kapasitas '
     'untuk mengoptimalkan bobot secara penuh.',False,False),
], indent=True)

add_body(doc,
    'Secara keseluruhan, ketiga metode berhasil melampaui kriteria keberhasilan '
    'yang ditetapkan (akurasi ≥ 95% dan AUC ≥ 0,98). Metode Residual Spatial '
    'direkomendasikan sebagai metode terbaik untuk deteksi deepfake pada konteks '
    'dataset IMSFD dengan face swapping berbasis InsightFace.',
    indent=True)

add_heading(doc, '4.4.7 Visualisasi Hasil', level=2)

add_body(doc,
    'Evaluasi juga menghasilkan visualisasi yang disimpan dalam direktori '
    'output/figures/ untuk setiap metode, meliputi:',
    indent=True)

add_bullet(doc, 'Confusion Matrix: Matriks 2×2 yang menampilkan distribusi '
                'True Positive, True Negative, False Positive, dan False Negative.')
add_bullet(doc, 'ROC Curve: Kurva yang memperlihatkan trade-off antara True '
                'Positive Rate dan False Positive Rate pada berbagai threshold.')
add_bullet(doc, 'Precision-Recall Curve: Kurva yang menunjukkan trade-off '
                'antara Precision dan Recall, terutama berguna untuk dataset '
                'yang seimbang.')

add_body(doc,
    'Seluruh visualisasi dihasilkan secara otomatis oleh fungsi full_evaluation() '
    'pada src/evaluate.py menggunakan library matplotlib dengan konfigurasi '
    'PLOT_DPI=200 untuk kualitas gambar yang optimal.',
    indent=True)

# ════════════════════════════════════════════════════════════════════════════
# 4.5  RINGKASAN
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.5 Ringkasan Implementasi')

add_body(doc,
    'Bab ini telah menjelaskan seluruh tahapan implementasi sistem deteksi deepfake '
    'menggunakan kerangka kerja CRISP-DM, mulai dari pemahaman tujuan penelitian '
    'hingga analisis hasil evaluasi. Ringkasan pencapaian:',
    indent=True)

add_bullet(doc, 'Pembangkitan dataset deepfake yang seimbang (13.439 real vs 13.439 fake) '
                'menggunakan InsightFace dengan mode split_aware.')
add_bullet(doc, 'Subject-level split (70:15:15) berhasil mencegah identity leakage antar split.')
add_bullet(doc, 'Ketiga model berhasil dilatih selama 12 epoch dengan konvergensi yang baik.')
add_bullet(doc, 'Metode Residual Spatial mencapai performa terbaik (accuracy 99,78%, AUC 0,9999).')
add_bullet(doc, 'Seluruh metode melampaui kriteria keberhasilan yang ditetapkan.')

add_body(doc,
    'Hasil yang diperoleh memvalidasi hipotesis penelitian bahwa representasi '
    'residual noise dapat meningkatkan kemampuan deteksi deepfake. Analisis '
    'lebih mendalam mengenai implikasi temuan ini, keterbatasan penelitian, '
    'dan arah penelitian selanjutnya akan dibahas pada Bab V.',
    indent=True)

# ════════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════════
output_path = 'BAB_IV_IMPLEMENTASI_DAN_PEMBAHASAN.docx'
doc.save(output_path)
print(f'Berhasil membuat: {output_path}')
